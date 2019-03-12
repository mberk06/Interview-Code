# Date: Summer 2018
# Author: Michael Berk
# Description: this file will be used to automate the Schork Report text, graphs, datess, and tables

#Classes:
#   Scrape: scrape commodity data from barchart.com
#   Read: Read data from excel files (this may be replaced with webscraping code)
#   Calculations and Helpers: perform necessary calculations
#   Write: write to a word document template

#imports
import selenium as sl
import time
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver import ActionChains
import selenium.webdriver.support.ui as ui
from selenium.webdriver.common.by import By

import xlrd
import csv
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.enum.style import WD_STYLE_TYPE

import numpy as np
from datetime import datetime, timedelta
import julian
import calendar

import random
import pandas as pd
from scipy.stats import norm
import statistics as stats
import math

from math import pi
from bokeh.plotting import figure, show, output_file
from bokeh.models import ColumnDataSource, Range1d, LabelSet, Label
from bokeh.io import export_png

#############################
#Scrape
#############################
class scrape():
    def __init__(self,downloadPath,u,p):
        #set up username and password
        self.username = u
        self.password = p

        #set download location
        options = webdriver.ChromeOptions()
        prefs = {"download.default_directory": downloadPath}
        options.add_experimental_option("prefs", prefs)
        self.d = webdriver.Chrome(chrome_options=options, executable_path=downloadPath+"/chromedriver.exe")

    #Purpose: create list of all urls
    #Parameters: list of all commodity, month letter, date
    #Return: full url list
    def createURL(self, com, month, year):
        URL = []
        c = 0
        for k,v in month.items():
            tempURL = "https://www.barchart.com/futures/quotes/"+com[c]+v+year[2:]+"/price-history/historical"
            URL.append(tempURL)
            c += 1
        return URL

    #Purpose: login to barchart
    #Parameters: username, password
    #Return: NA
    def login(self, u, p):
        self.d.find_element_by_partial_link_text("Log In or Sign Up").click()
        self.d.find_element_by_name("email").send_keys(u)
        self.d.find_element_by_name("password").send_keys(p)
        self.d.find_element_by_xpath("//*[contains(text(), 'Log In')]").click()
        time.sleep(3)

    #Purpose: downlaod the commodity
    #Parameters: list of URLs
    #Return: NA
    def downloadCommodity(self, URL):
        #iterate through URLs
        for u in range(0,len(URL)):
            print("Scraping from:",URL[u])

            #go to URL and download data
            self.d.get(URL[u])
            if u == 0:
                self.login(self.username, self.password)

            time.sleep(3)
            WebDriverWait(self.d, 9).until(EC.element_to_be_clickable((By.PARTIAL_LINK_TEXT, "max")))
            self.d.find_element_by_partial_link_text("max").click()
            time.sleep(3)

            print("Success")

        #close the driver
        self.d.close()

    """
    #Purpose: scrape expiration date
    #Parameters: commodity type
    #Return: [future expiration, option expiration]
    def getExpiration(self, com):
        # flow control for commodities
        if com == "NG":
            urls = ["https://www.cmegroup.com/trading/energy/natural-gas/natural-gas_product_calendar_futures.html?optionProductId=191&optionExpiration=191-Q8",
                    "https://www.cmegroup.com/trading/energy/natural-gas/natural-gas_product_calendar_options.html?optionProductId=191&optionExpiration=191-Q8"]
            return(self.readExpirationTable(urls))
        elif com == "RBOB":
            urls = ["https://www.cmegroup.com/trading/energy/refined-products/rbob-gasoline_product_calendar_futures.html?optionProductId=195",
                    "https://www.cmegroup.com/trading/energy/refined-products/rbob-gasoline_product_calendar_options.html?optionProductId=195"]
            return(self.readExpirationTable(urls))
        elif com == "WTI":
            urls = ["https://www.cmegroup.com/trading/energy/crude-oil/light-sweet-crude_product_calendar_futures.html",
                    "https://www.cmegroup.com/trading/energy/crude-oil/light-sweet-crude_product_calendar_options.html"]
            return(self.readExpirationTable(urls))
        elif com == "ULSD":
            urls = ["https://www.cmegroup.com/trading/energy/refined-products/heating-oil_product_calendar_futures.html",
                    "https://www.cmegroup.com/trading/energy/refined-products/heating-oil_product_calendar_options.html#optionProductId=194"]
            return(self.readExpirationTable(urls))
        elif com == "Brent":
            urls = ["https://www.cmegroup.com/trading/energy/crude-oil/brent-ice-calendar-swap-futures_product_calendar_futures.html",
                    "https://www.cmegroup.com/trading/energy/crude-oil/brent-ice-calendar-swap-futures_product_calendar_options.html"]
            return(self.readExpirationTable(urls))
        elif com == "Gasoil":
            urls = ["https://www.cmegroup.com/trading/energy/refined-products/european-gasoil-ice-futures_product_calendar_futures.html",
                    "https://www.cmegroup.com/trading/energy/refined-products/european-gasoil-ice-futures_product_calendar_options.html#optionProductId=5790"]
            return(self.readExpirationTable(urls))
        else:
            print("Incorrect com specified")
            return None

    #Purpose: read table
    #Parameters: url array
    #Return: [future expiration, option expiration]
    def readExpirationTable(self, urls):
        #read future and option
        self.d.get(urls[0])
        time.sleep(2)
        future = self.d.find_element_by_xpath("/html[1]/body[1]/div[2]/div[2]/div[3]/div[1]/div[1]/div[2]/div[2]/div[3]/div[4]/table[1]/tbody[1]/tr[1]/td[2]").text.split()
        dateF = future[3]+" "+future[4]+" "+future[5]
        self.d.get(urls[1])
        time.sleep(2)
        option = self.d.find_element_by_xpath("/html[1]/body[1]/div[2]/div[2]/div[3]/div[1]/div[1]/div[2]/div[2]/div[3]/div[4]/table[1]/tbody[1]/tr[1]/td[2]").text.split()
        dateO = option[3]+" "+option[4]+" "+option[5]

        return [dateF, dateO]
    """

#############################
#Read
#############################
class read():
    def __init__(self, excelPath, csvPath):
        self.excelPath = excelPath
        self.csvPath = csvPath

    #Purpose: get data from excel file
    #Parameters: type of commodity to get
    #Return: list with data stored
    def getDataExcel(self, commodity):
        print("Reading:", commodity)

        #initialize data arrays
        dateArr = np.array([0])
        openArr = np.array([0])
        highArr = np.array([0])
        lowArr = np.array([0])
        settleArr = np.array([0])
        changeArr = np.array([0])

        #open the workbook
        workbook = xlrd.open_workbook(self.excelPath)

        #flow control for commodities
        if commodity == "NG":
            return self.getColumns(workbook.sheet_by_name("NYMEX"), [0,1,2,3,4,5])
        elif commodity == "RBOB":
            return self.getColumns(workbook.sheet_by_name("NYMEX"), [16,17,18,19,20,21])
        elif commodity == "WTI":
            return self.getColumns(workbook.sheet_by_name("NYMEX"), [10,11,12,13,14,15])
        elif commodity == "ULSD":
            return self.getColumns(workbook.sheet_by_name("NYMEX"), [22,23,24,25,26,27])
        elif commodity == "Brent":
            return self.getColumns(workbook.sheet_by_name("ICE"), [0,1,2,3,4,5])
        elif commodity == "Gasoil":
            return self.getColumns(workbook.sheet_by_name("ICE"), [6,7,8,9,10,11])
        else:
            print("Incorrect commodity specified")
            return None

    #Purpose: get 4 columns from excel file
    #Parameters: worksheet to read from, 4 columns to get
    #Return: list with data stored
    def getColumns(self, worksheet, cols):
        #initialize data arrays
        dateArr = np.array([0])
        openArr = np.array([0])
        highArr = np.array([0])
        lowArr = np.array([0])
        settleArr = np.array([0])

        for row in range(worksheet.nrows):
            #append values to array
            dateArr = np.append(dateArr, worksheet.cell(row, cols[0]).value)
            openArr = np.append(openArr, worksheet.cell(row, cols[1]).value)
            highArr = np.append(highArr, worksheet.cell(row, cols[2]).value)
            lowArr = np.append(lowArr, worksheet.cell(row, cols[3]).value)
            settleArr = np.append(settleArr, worksheet.cell(row, cols[4]).value)

        return [dateArr[6:].astype(str), openArr[6:].astype(np.float), highArr[6:].astype(np.float), lowArr[6:].astype(np.float), settleArr[6:].astype(np.float)]

    #Purpose: read from csv (web scraped)
    #Parameters: commodity type, month letter
    #Return: list with data stored
    def getDataCSV(self, com, month):
        #get today's date
        date = datetime.now().strftime("%m-%d-%Y")

        #flow control for commodities
        if com == "NG":
            return self.readCSV(self.csvPath+"ng"+month.lower()+date[-2:]+"_price-history-"+date+".csv")
        elif com == "RBOB":
            return self.readCSV(self.csvPath+"rb"+month.lower()+date[-2:]+"_price-history-"+date+".csv")
        elif com == "WTI":
            return self.readCSV(self.csvPath+"cl"+month.lower()+date[-2:]+"_price-history-"+date+".csv")
        elif com == "ULSD":
            return self.readCSV(self.csvPath+"ho"+month.lower()+date[-2:]+"_price-history-"+date+".csv")
        elif com == "Brent":
            return self.readCSV(self.csvPath+"cb"+month.lower()+date[-2:]+"_price-history-"+date+".csv")
        elif com == "Gasoil":
            return self.readCSV(self.csvPath+"lf"+month.lower()+date[-2:]+"_price-history-"+date+".csv")
        else:
            print("Incorrect com specified")
            return None

    #Purpose: read data from csv path
    #Parameters: path
    #Return: list with data stored
    def readCSV(self, path):
        csvfile = open(path, 'r')
        #date, open, high, low, close, change
        dateArr = np.array([0])
        openArr = np.array([0])
        highArr = np.array([0])
        lowArr = np.array([0])
        settleArr = np.array([0])
        changeArr = np.array([0])

        index = 10000

        for i, line in enumerate(csvfile):
            #avoid first and last line
            if not i == 0 and not i == 500 and "Barchart.com" not in line:
                #add data to arrays
                array = line.split(',')
                dateArr = np.append(dateArr, array[0])
                openArr = np.append(openArr, array[1])
                highArr = np.append(highArr, array[2])
                lowArr = np.append(lowArr, array[3])
                settleArr = np.append(settleArr, array[4])
                changeArr = np.append(changeArr, array[5])

        return self.correctIndex([dateArr[1:].astype(str), openArr[1:].astype(np.float), highArr[1:].astype(np.float), lowArr[1:].astype(np.float), settleArr[1:].astype(np.float), changeArr[1:].astype(np.float)])

    #Purpose: subset data so that 3 rows in a row do not have the same values (within the row)
    #Parameters:
    #Return: subsetted data
    def correctIndex(self, data):
        #initialize
        twoRowsBackIsSame = False
        priorRowIsSame = False

        #iterate through data
        for r in range(0, len(data[0])):
            #get row data
            open = data[1][r]
            high = data[2][r]
            low = data[3][r]
            close = data[4][r]

            #flow control to get index
            currentRowIsSame = open == high and high == low and low == close

            #3 rows in a row are the same
            if twoRowsBackIsSame and priorRowIsSame and currentRowIsSame:
                returnList = []
                for d in data:
                    returnList.append(d[0:r-2])
                return returnList
            #2 rows in a row are the same
            elif currentRowIsSame and priorRowIsSame:
                twoRowsBackIsSame = True
            #1 row is same
            elif currentRowIsSame:
                priorRowIsSame = True
            #no rows are the same
            else:
                priorRowIsSame = False
                twoRowsBackIsSame = False
        return data

#############################
#Calculations and Helpers
#############################

class calculations():
    #Purpose: get daily/weekly/monthly levels of support and resistance
    #Parameters: specify daily/weekly/monthly, specify tau, data from read.getData("Commodity")
    #Return: dictionary with the levels of support and resistance
    def supportAndResistance(self, timeFrame, tau, data):
        print("Calculating",timeFrame,"levels of support and resistance for",tau,"levels")

        #create data
        high = np.array([])
        low = np.array([])
        settleVal = data[4][0]

        #calculate log values for high and low prices
        if timeFrame == "Daily":
            daysBack = 21
            for i in range(len(data[0]) - daysBack, len(data[0])):
                high = np.append(high, math.log(data[2][i]/data[2][i-1]))
                low = np.append(low, math.log(data[3][i]/data[3][i-1]))
        elif timeFrame == "Weekly":
            daysBack = 63
            for i in range(len(data[0]) - daysBack, len(data[0])):
                high = np.append(high, math.log(data[2][i]/data[2][i-1]))
                low = np.append(low, math.log(data[3][i]/data[3][i-1]))
        elif timeFrame == "Monthly":
            if len(data[0]) < 360:
                daysBack = len(data[0])
                print("Not using 360 days to calculate monthly levels, instead using: "+str(daysBack))
            else:
                daysBack = 360
            for i in range(len(data[0]) - daysBack, len(data[0])):
                high = np.append(high, math.log(data[2][i]/data[2][i-1]))
                low = np.append(low, math.log(data[3][i]/data[3][i-1]))
        else:
            print("Incorrect time frame specified")
            return None

        #set up return data
        returnDict = {}
        for k in tau:
            returnDict[k] = -1

        #iterate through support and resistance data
        for d in [high,low]:
            #initial calculations
            mean = np.mean(d)
            sd = np.std(d)
            var = sd ** 2
            drift = mean + (var / 2)

            #specify tau for level of support and resistance
            if (d == high).all():
                tempTau = tau[0:round(len(tau)/2)]
            else:
                tempTau = tau[round(len(tau)/2):len(tau)]

            #run 10000 times to get percentile from values
            tempArr = np.array([])
            for i in range(0, 10000):
                #perform calculation
                value = settleVal * math.exp(drift + sd * norm.ppf(random.uniform(0, 1)))
                tempArr = np.append(tempArr, value)

            #get quantiles to use
            for t in tempTau:
                #get percentile that corresponds with tau
                returnDict[t] = np.percentile(tempArr, t*100)

        return returnDict

    #Purpose: create a paragraph for a commodity
    #Parameters: commodity type, levels of support and resistance
    #Return: paragraph string
    def createParagraphs(self, commodityName, data, quantiles):
        #get peak, bottom, close
        peak = data[2][0]
        bottom = data[3][0]
        close = data[4][0]
        change = data[5][0]

        #get sentences
        opener = self.getMonth()[0] + " " + commodityName
        peakString = self.createSentence("peaked", peak, quantiles)
        bottomString = self.createSentence("bottomed", bottom, quantiles)
        closeString = self.createSentence("closed", close, quantiles)

        paragraph1 = opener + " " + peakString + ", " + bottomString + ", and " + closeString + "."

        paragraph2 = "As far as today goes for " + self.getMonth()[1] + " " + commodityName + ", a drop below " + str(
            round(quantiles[0.37], 3)) + \
             " alerts us to weakness towards our " + str(
            round(quantiles[0.25], 3)) + " second support point. Below here, we look " \
                                            "for support at our " + str(
            round(quantiles[0.05], 3)) + " third level of support. Then again, strength above " + \
             str(round(quantiles[0.63], 3)) + " opens the door to our " + str(
            round(quantiles[0.75], 3)) + " second level of resistance. " \
                                            "Through here, we will look for resistance to hold at " + str(
            round(quantiles[0.95], 3)) + "."

        return[paragraph1,paragraph2]

    #Purpose: create the sentences that make up paragraphs
    #Parameters: peak/bottom/close, value, closest value
    #Return: string of sentence that will be added to paragraph
    def createSentence(self, type, value, quantiles):
        #perform initial calcualations
        closest = self.closest(value, quantiles.values())
        ticks = round(closest - value, 3) * 1000

        #get corresponding tau
        level = None
        for key, val in quantiles.items():
            if val == closest:
                level = self.getLevel(key)

        #check if value is above or below closest value
        aboveOrBelow = "below"
        if ticks > 0:
            aboveOrBelow = "above"
        ticks = abs(ticks)

        return type + " " + str(ticks) + " ticks " + aboveOrBelow + " our " + level

    #Purpose: find value in array that is closes to val
    #Parameters: value to test proximity, array of values
    #Return: values closest to val
    def closest(self, val, arr):
        closest = -1
        dif = 10000
        for a in arr:
            if abs(val - a) < dif:
                dif = abs(val - a)
                closest = a
        return closest

    #Purpose: get today's month and month letter
    #Parameters: NA
    #Return: today's month, today's month letter, next month's letter
    def getMonth(self):
        current_month = datetime.now().strftime('%m')
        monthDict = {"01": ["January","F","G"], "02": ["Febrary","G","H"], "03": ["March","H","J"], "04": ["April","J","K"], "05": ["May","K","M"],
                     "06": ["June","M","N"], "07": ["July","N","Q"], "08": ["August","Q","U"], "09": ["September","U","V"],"10": ["October","V","X"],
                     "11": ["November","X","Z"], "12": ["December","Z","F"]}
        return monthDict[current_month]

    #Purpose: get today's full date
    #Parameters: NA
    #Return: today's full date
    def getFullDate(self):
        now = datetime.now()
        return [now.strftime("%m/%d/%Y").replace(" 0", " "),now.strftime("%m-%d-%Y").replace(" 0", " ")]

    #Purpose: get today's year
    #Parameters: NA
    #Return: today's year
    def getYear(self):
        now = datetime.now()
        return now.strftime("%Y")

    # Purpose: get Contract
    # Parameters: NA
    # Return: dict of contracts with replacement
    def getContract(self, month):
        return {"{1Cont}": "NYMEX NG" + month + self.getYear()[2:],"{2Cont}": "NYMEX CL" + month + self.getYear()[2:],
                "{3Cont}": "ICE CB" + month + self.getYear()[2:],"{4Cont}": "NYMEX RB" + month + self.getYear()[2:],
                "{5Cont}": "NYMEX HO" + month + self.getYear()[2:],"{6Cont}": "ICE LF" + month + self.getYear()[2:]}

    #Purpose: get text from levels
    #Parameters: value
    #Return: string corresponding to level of support or reistance
    def getLevel(self, tau):
        dict = {0.95: "third level of resistance",
                0.75: "second level of resistance",
                0.63: "first level of resistance",
                0.37: "first level of support",
                0.25: "second level of support",
                0.05: "third level of support"}
        return dict[tau]

    #Purpose: get section header values
    #Parameters: commodity, data
    #Return: list of commodity header values
    def getHeaderValues(self, com, data):
        high = data[2][0]
        low = data[3][0]
        close = data[4][0]
        change = data[5][0]
        return [high, low, close, change]

    #Purpose: get trend as of date
    #Parameters: commodity list, data
    #Return: list of trend
    def getTrend(self, com, data):
        #inial setup
        close = data[4]
        dates = data[0]
        today = datetime(datetime.today().year, datetime.today().month, datetime.today().day).strftime('%m/%d/%y')
        todayDate = datetime.today()
        dayOfWeek = datetime.today().weekday()
        returnArr = []

        #daily
        if dayOfWeek == 0:
            returnArr.append(str(close[(3)]))
        elif dayOfWeek == 6:
            returnArr.append(str(close[(2)]))
        else:
            returnArr.append(str(close[(1)]))

        #weekly
        if dayOfWeek > 4:
            returnArr.append(str(close[(dayOfWeek-4)]))
        else:
            returnArr.append(str(close[(dayOfWeek+3)]))

        #monthly
        first = todayDate.replace(day=1)
        priorMonthLastDay = first - timedelta(days=1)
        comNum = -1
        if com == "NG":
            comNum = 1
            index = np.where(dates == (self.weekdayTest(priorMonthLastDay-timedelta(days=3)).strftime('%m/%d/%y')))[0][0]
            returnArr.append(str(close[index]))
        elif com == "WTI":
            comNum = 2
            index = np.where(dates == (self.weekdayTest(priorMonthLastDay.replace(day=25)-timedelta(days=3)).strftime('%m/%d/%y')))[0][0]
            returnArr.append(str(close[index]))
        elif com == "Brent":
            comNum = 3
            index = np.where(dates == (self.weekdayTest(priorMonthLastDay.replace(day=1)-timedelta(days=1)).strftime('%m/%d/%y')))[0][0]
            returnArr.append(str(close[index]))
        elif com == "RBOB":
            comNum = 4
            index = np.where(dates == (self.weekdayTest(priorMonthLastDay)).strftime('%m/%d/%y'))[0][0]
            returnArr.append(str(close[index]))
        elif com == "ULSD":
            comNum = 5
            index = np.where(dates == (self.weekdayTest(priorMonthLastDay)).strftime('%m/%d/%y'))[0][0]
            returnArr.append(str(close[index]))
        elif com == "Gasoil":
            comNum = 6
            index = np.where(dates == (self.weekdayTest(priorMonthLastDay.replace(day=14)-timedelta(days=3)).strftime('%m/%d/%y')))[0][0]
            returnArr.append(str(close[index]))
        else:
            print("Incorrect com specified")
            return None

        return returnArr

    #Purpose: test if weekday
    #Parameters: day
    #Return: nearest weekday in the past
    def weekdayTest(self, day):
        bdays = pd.to_datetime(pd.bdate_range(datetime.today()-timedelta(days=260),datetime.today()))
        for i in range(0,360):
            date = datetime(year=day.year, month=day.month, day=day.day)
            if date in bdays:
                return day
            else:
                day = day-timedelta(days=1)
        print("weekdayTest error")
        return None

    #Purpose: get all text in tables for all commdities, tau, [daily, weekly, monthly], [high, low, close, change], [dayTrend, weekTrend, monthTrend]
    #Parameters: commodity
    #Return: all text to be replaced for each
    def getCommodityTables(self, com, tau, levels, header, trend):
        #intial setup of words to replace
        trendWords = ["{TrendDay}","{TrendWeek}","{TrendMonth}"]
        headerWords = ["{High}","{Low}","{Close}","{Change}"]
        dailyWords = ["{Res3Day}","{Res2Day}","{Res1Day}","{Sup1Day}","{Sup2Day}","{Sup3Day}"]
        weeklyWords = ["{Res3Week}","{Res2Week}","{Res1Week}","{Sup1Week}","{Sup2Week}","{Sup3Week}"]
        monthlyWords = ["{Res3Month}","{Res2Month}","{Res1Month}","{Sup1Month}","{Sup2Month}","{Sup3Month}"]
        returnDict = {}
        comNumber = -1

        #get commodity number
        if com == "NG":
            comNumber = 1
        elif com == "RBOB":
            comNumber = 2
        elif com == "WTI":
            comNumber = 3
        elif com == "ULSD":
            comNumber = 4
        elif com == "Brent":
            comNumber = 5
        elif com == "Gasoil":
            comNumber = 6
        else:
            print("Incorrect com specified")
            return None

        #iterate through trend words
        for i in range(0,len(trendWords)):
            trendWord = trendWords[i][:1] + str(comNumber) + trendWords[i][1:]
            returnDict[trendWord] = trend[i]

        #iterate through header words
        for i in range(0,len(headerWords)):
            headerWord = headerWords[i][:1] + str(comNumber) + headerWords[i][1:]
            returnDict[headerWord] = str(round(header[i],3))

        #iterate through support and resistance words
        for i in range(0, len(levels[0])):
            #change words to be replaced based on commodity
            dailyWord = dailyWords[i][:1] + str(comNumber) + dailyWords[i][1:]
            weeklyWord = weeklyWords[i][:1] + str(comNumber) + weeklyWords[i][1:]
            monthlyWord = monthlyWords[i][:1] + str(comNumber) + monthlyWords[i][1:]

            #edit return dictionary
            returnDict[dailyWord] = str(round(levels[0][tau[i]],3))
            returnDict[weeklyWord] = str(round(levels[1][tau[i]],3))
            returnDict[monthlyWord] = str(round(levels[2][tau[i]],3))

        return returnDict

    #Purpose: create price graphs for each commodity
    #Parameters: commodity, data
    #Return: price graph
    def createGraph(self, title, data, levels, types, date, path):
        #create data frame
        graphDict = {"date":data[0][:75],"open":data[1][:75],"high":data[2][:75],"low":data[3][:75],"close":data[4][:75]}
        df = pd.DataFrame(graphDict)
        df["date"] = pd.to_datetime(df["date"])

        inc = df.close > df.open
        dec = df.open > df.close
        w = 12 * 60 * 60 * 1000  # half day in ms

        p = figure(x_axis_type="datetime", plot_width=1080, plot_height=1100, title=title, toolbar_location=None)
        p.xaxis.major_label_orientation = pi / 4
        p.grid.grid_line_alpha = 0.3

        p.segment(df.date, df.high, df.date, df.low, color="black")
        p.vbar(df.date[inc], w, df.open[inc], df.close[inc], fill_color="white", line_color="black")
        p.vbar(df.date[dec], w, df.open[dec], df.close[dec], fill_color="black", line_color="black")

        #set axis titles and font size
        p.yaxis.axis_label = ""

        p.yaxis.axis_label_text_font = "verdana"
        p.yaxis.major_label_text_font = "verdana"
        p.xaxis.major_label_text_font = "verdana"
        p.title.text_font = "verdana"
        p.title.text_font_size = "25pt"

        for i in range(0, len(types)):
            #initial setup
            type = types[i]
            level = levels[i]

            #Graph levels of support and resistance
            if type == "Daily" or type == "Weekly" or type == "Monthly":
                #initial setup
                levelsDict = {0.95:"3",0.75:"2",0.63:"1",0.37:"1",0.25:"2",0.05:"3"}
                if type == "Daily":
                    divNum = 6
                elif type == "Weekly":
                    divNum = 4
                else:
                    divNum = 2.9

                #iterate through levels
                for k,v in level.items():
                    #add line and label
                    if k > 0.5:
                        p.step([df["date"][0],df["date"][round(len(df["date"])/divNum)]],[v,v], line_width=1, line_color="green")
                        label = Label(x=df["date"][0], y=v, text=levelsDict[k]+type[0], render_mode='css', text_color = "green",
                                      background_fill_alpha=0, text_font_size="9pt", x_offset = 15, y_offset = -7.5)
                    else:
                        p.step([df["date"][0],df["date"][round(len(df["date"])/divNum)]],[v,v], line_width=1, line_color="red")
                        label = Label(x=df["date"][0], y=v, text=levelsDict[k]+type[0], render_mode='css', text_color = "red",
                              background_fill_alpha=0, text_font_size="9pt", x_offset = 15, y_offset = -7.5)

                    p.add_layout(label)
            else:
                print("Incorrect type of level specified, did not add level to graph")

        export_png(p, filename=path+title+date+".png")

        return path+title+date+".png"

#############################
#Write
#############################
class write():
    #initialize file paths
    def __init__(self, path, savePath):
        self.savePath = savePath
        self.doc = Document(path)
        self.style = self.doc.styles['Normal']
        font = self.style.font
        font.name = 'Verdana'

        #make small style
        styles = self.doc.styles
        self.small = styles.add_style('Small', WD_STYLE_TYPE.PARAGRAPH)
        font = self.small.font
        font.size = Pt(7.5)
        font.name = 'Verdana'

    #Purpose: replace a word and save the document to save doc path
    #Params: dictionary of words to be replaced with their replacemenet
    #Return: NA
    def replaceWord(self, words):
        for (k,v) in words.items():
            # replace paragraphs
            for paragraph in self.doc.paragraphs:
                if k in paragraph.text:
                    text = paragraph.text
                    if "Image" in k:
                        paragraph.text = text.replace(k,"")
                        r = paragraph.add_run()
                        r.add_picture(v, width=Inches(3.5))
                    else:
                        paragraph.text = text.replace(k, v)
                    paragraph.style = self.style
            # replace table paragraphs
            for t in self.doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if k in paragraph.text:
                                text = paragraph.text
                                if "Image" in k:
                                    paragraph.text = text.replace(k,"")
                                    r = paragraph.add_run()
                                    r.add_picture(v,width=Inches(3.5))
                                else:
                                    paragraph.text = text.replace(k, v)
                                paragraph.style = self.small

        self.doc.save(self.savePath)

#############################
#Run
#############################
#specify paths and other initial setup
excelPath = "C:/Users/Owner/Desktop/Automation/ScrapedData/"
csvPath = "C:/Users/Owner/Desktop/Automation/ScrapedData/"
docPath = "C:/Users/Owner/Desktop/Automation/Word Files/template.docx"
saveDocPath = "C:/Users/Owner/Desktop/Automation/Word Files/save.docx"
graphPath = "C:/Users/Owner/Desktop/Automation/Graphs/"

username = "NA"
password = "NA"

shouldScrape = False #should scrape data?
dailyLevels = True #include daily levels?
weeklyLevels = True #include weekly levels?
monthlyLevels = True #include monthly levels?
comMonth = {"Natural Gas":"U", "Oil":"U", "ICE Brent":"U", "Heating Oil":"U", "Middle Distillates":"U", "Gasoil":"U"} #specify month for commodity

#initialize classes
if shouldScrape:
    scrape = scrape(csvPath,username,password)
read = read(excelPath,csvPath)
calculations = calculations()
write = write(docPath, saveDocPath)

#extra setup
comCodes = ["NG","CL","CB","RB","HO","LF"]
commodities = ["NG", "WTI", "Brent", "RBOB", "ULSD", "Gasoil"]
names = ["Natural Gas", "Oil", "ICE Brent", "Heating Oil", "Middle Distillates", "Gasoil"]
graphTitles = ["NYMEX NG Price", "NYMEX WTI Price", "ICE Brent Price", "NYMEX RBOB Price", "NYMEX ULSD Price", "ICE Gasoil Price"]
tau = [0.95,0.75,0.63,0.37,0.25,0.05]

#scrape data
if shouldScrape:
    commodityURLs = scrape.createURL(comCodes,comMonth,calculations.getYear())
    scrape.downloadCommodity(commodityURLs)
    print("Scraping Complete")

#AUTOMATE TEXT
#get words to replace
words = {}
for i in range(0, len(names)):
    print("Working on: ", names[i])
    #read data
    data = read.getDataCSV(commodities[i],comMonth[names[i]])

    #get quantiles
    header = calculations.getHeaderValues(commodities[i], data)
    trend = calculations.getTrend(commodities[i], data)
    daily = calculations.supportAndResistance("Daily",tau,data)
    weekly = calculations.supportAndResistance("Weekly",tau,data)
    monthly = calculations.supportAndResistance("Monthly",tau,data)

    #create paragraphs
    paragraphs = calculations.createParagraphs(names[i], data, daily)

    #create words to replace
    words["{"+str(i+1)+"Paragraph1}"] = paragraphs[0]
    words["{"+str(i+1)+"Paragraph2}"] = paragraphs[1]

    #automate tables
    tableWords = calculations.getCommodityTables(commodities[i],tau,[daily,weekly,monthly],header,trend)

    #add values to words
    words.update(tableWords)
    words.update(calculations.getContract(comMonth["Natural Gas"]))

    #graph
    levelsArr = []
    levelsNames = []
    if dailyLevels:
        levelsArr.append(daily)
        levelsNames.append("Daily")
    if weeklyLevels:
        levelsArr.append(weekly)
        levelsNames.append("Weekly")
    if monthlyLevels:
        levelsArr.append(monthly)
        levelsNames.append("Monthly")

    words["{"+str(i+1)+"Image}"] = calculations.createGraph(graphTitles[i],data,levelsArr,levelsNames,calculations.getFullDate()[1],graphPath)

#AUTOMATE DATES
#add words to replace
words["{Year}"] = calculations.getYear()
words["{Month}"] = calculations.getMonth()[0]

#run replacements
write.replaceWord(words)
